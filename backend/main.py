from datetime import datetime
from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from auth import Auth, LoginBody
from db import init_db
from models import Printer, PrintJob, Document, Log, Status, SystemConfig
from threading import Thread
from time import sleep
import docx
import io
import PyPDF2

app = FastAPI()
auth = Auth()
db = init_db()

def log_callback(printer: Printer, print_job: PrintJob):
    assert (
        printer.id is not None
        and print_job.id is not None
        and print_job.student_id is not None
    )
    log = Log(
        printer_id=printer.id,
        print_job=print_job,
        student_id=print_job.student_id,
        description=f"[END] user={print_job.student_id} doc={print_job.document.file_name}",
        date=datetime.now(),
    )

    # this is thread safe
    db.add_log(log)


for printer in db.get_printers():
    printer_thread = Thread(target=printer.run, args=(log_callback,))
    printer_thread.start()

def add_pages_each_sem():
    time = datetime.now()
    # first semester starts from 15th August
    # second semester starts from 1st January
    while True:
        if time.month < 8 or (time.month == 8 and time.day < 15):
            next_sem = datetime(time.year, 8, 15)
        else:
            next_sem = datetime(time.year + 1, 1, 1)
        delta = next_sem - time

        sleep(delta.total_seconds())
        db.add_pages_to_students()
        # sleep for 1 day to avoid multiple calls
        sleep(24 * 60 * 60 + 1)

Thread(target=add_pages_each_sem).start()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def authorize(request: Request, call_next):
    if request.url.path in ["/login", "/docs", "/openapi.json"]:
        return await call_next(request)
    token = request.headers.get("Authorization") or request.cookies.get("token")
    if token is None:
        return Response(status_code=401, content="Unauthorized")

    user_id = auth.is_authorised(token)
    if user_id is None:
        return Response(status_code=401, content="Unauthorized")
    request.state.user = user_id
    response = await call_next(request)
    return response


@app.post("/login")
async def login(response: Response, login_data: LoginBody):
    user, token = auth.login(login_data.login, login_data.password)
    if user is None or token is None:
        return {"error": "Invalid login or password"}
    response.set_cookie("token", token, httponly=True)
    return {"token": token, "role": auth.role(user.user_id)}


@app.get("/printer")
async def get_printers():
    return db.get_printers()


@app.post("/printer/add")
async def add_printer(request: Request, printer: Printer):
    if auth.role(request.state.user) != "spso":
        raise HTTPException(status_code=403, detail="Forbidden")
    printers = db.get_printers()
    assert len(printers) > 0 and printers[-1].id is not None
    next_printer_id = printers[-1].id + 1
    printer.id = next_printer_id
    db.add_printer(printer)
    if printer.status == Status.ENABLED:
        Thread(target=printer.run, args=(log_callback,)).start()
    return {"id": printer.id}


@app.post("/printer/update")
async def update_printer(request: Request, printer: Printer):
    if auth.role(request.state.user) != "spso":
        raise HTTPException(status_code=403, detail="Forbidden")
    if printer.id is None:
        raise HTTPException(status_code=400, detail="Printer ID is required")
    try:
        p = db.update_printer(printer)
        return {"id": p.id}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/printer/delete/{printer_id}")
async def delete_printer(request: Request, printer_id: int):
    if auth.role(request.state.user) != "spso":
        raise HTTPException(status_code=403, detail="Forbidden")
    printer = db.get_printer_by_id(printer_id)
    if printer is None:
        raise HTTPException(status_code=404, detail="Printer not found")
    if printer.status == Status.ENABLED:
        raise HTTPException(status_code=400, detail="Printer is enabled. Please disable it first")
    try:
        db.delete_printer(printer_id)
        return {"id": printer_id}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/printer/print")
async def add_printjob(request: Request, print_job: PrintJob):
    if auth.role(request.state.user) != "student":
        raise HTTPException(status_code=403, detail="Forbidden")

    printers = db.get_printers()
    least_busy_printer = min(printers, key=lambda printer: len(printer._printing_queue))
    assert least_busy_printer is not None and least_busy_printer.id is not None

    # 32 bit for timestamp, 32 bit for printer id
    print_job.id = (
        int(datetime.now().timestamp() * 1000) << 32
    ) | least_busy_printer.id
    least_busy_printer.add_printjob(print_job)
    print_job.student_id = request.state.user

    log = Log(
        description=f"[QUEUE] user={print_job.student_id} doc={print_job.document.file_name}",
        student_id=request.state.user,
        printer_id=least_busy_printer.id,
        print_job=print_job,
        date=datetime.now(),
    )
    db.add_log(log)

    return {"printer_id": least_busy_printer.id, "print_job_id": print_job.id}

@app.post("/printer/{printer_id}/print")
async def add_printjob_to_printer(request: Request, printer_id: int, print_job: PrintJob):
    if auth.role(request.state.user) != "student":
        raise HTTPException(status_code=403, detail="Forbidden")
    printer = db.get_printer_by_id(printer_id)
    if printer is None:
        raise HTTPException(status_code=404, detail="Printer not found")
    assert printer.id is not None
    print_job.student_id = request.state.user
    print_job.id = (
        int(datetime.now().timestamp() * 1000) << 32
    ) | printer.id
    printer.add_printjob(print_job)

    log = Log(
        description=f"[QUEUE] user={print_job.student_id} doc={print_job.document.file_name}",
        student_id=request.state.user,
        printer_id=printer.id,
        print_job=print_job,
        date=datetime.now(),
    )
    db.add_log(log)

    return {"printer_id": printer.id, "print_job_id": print_job.id}


@app.get("/printjob")
async def get_printjobs(request: Request):
    printers = db.get_printers()
    printjobs = [
        printjob for printer in printers for printjob in printer._printing_queue
    ]
    if auth.role(request.state.user) == "student":
        printjobs = [
            printjob
            for printjob in printjobs
            if printjob.student_id == request.state.user
        ]
    return printjobs


@app.post("/upload")
async def upload_file(request: Request):
    if auth.role(request.state.user) != "student":
        raise HTTPException(status_code=403, detail="Forbidden")
    student = db.get_student_by_id(request.state.user)
    assert student is not None
    form = await request.form()
    file = form.get("file")
    if file is None:
        raise HTTPException(status_code=400, detail="File is required")
    if isinstance(file, str):
        raise HTTPException(status_code=400, detail="Invalid file")

    file_name = file.filename
    if file_name is None:
        raise HTTPException(status_code=400, detail="Invalid file name")

    true_name, ext = file_name.split(".")
    if ext not in db.get_system_config().allowed_file_types:
        raise HTTPException(status_code=400, detail="Invalid file type")
    file_bytes = await file.read()
    await file.close()

    match ext:
        # case docx and doc
        case "docx", "doc":
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
            except Exception:
                raise HTTPException(status_code=400, detail="Invalid DOCX file")
            page_cnt = sum(p.contains_page_break for p in doc.paragraphs) + 1
            if student.remaining_pages < page_cnt:
                raise HTTPException(status_code=400, detail="Insufficient pages")

            file_id = db.upload_file(file_bytes, ext)
            document = Document(
                file_id=file_id,
                file_name=true_name,
                file_type=ext,
                pages=page_cnt,
            )
            db.add_document(document)
        case "pdf":
            try:
                pdf = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            except Exception as e:
                raise HTTPException(status_code=400, detail="Invalid PDF file")
            page_cnt = len(pdf.pages)
            if student.remaining_pages < page_cnt:
                raise HTTPException(status_code=400, detail="Insufficient pages")
            file_id = db.upload_file(file_bytes, ext)
            document = Document(
                file_id=file_id,
                file_name=true_name,
                file_type=ext,
                pages=page_cnt,
            )
            db.add_document(document)

        case _:
            raise HTTPException(status_code=400, detail="Invalid file type")

    return document

@app.get("/student")
async def get_student(request: Request):
    if auth.role(request.state.user) != "student":
        raise HTTPException(status_code=403, detail="Forbidden")
    student_id = request.state.user
    student = db.get_student_by_id(student_id)
    assert student is not None
    return student

@app.get("/student/{student_id}")
async def get_student_by_id(request: Request, student_id: int):
    if auth.role(request.state.user) != "spso":
        raise HTTPException(status_code=403, detail="Forbidden")
    student = db.get_student_by_id(student_id)
    if student is None:
        raise HTTPException(status_code=404, detail="Student not found")
    return student

@app.get("/system")
async def get_system_config():
    return db.get_system_config()

@app.post("/system/update")
async def update_system_config(request: Request, config: SystemConfig):
    if auth.role(request.state.user) != "spso":
        raise HTTPException(status_code=403, detail="Forbidden")
    db.update_system_config(config)
    return config

@app.get("/log")
async def get_logs(request: Request):
    logs = db.get_logs()
    if auth.role(request.state.user) == "student":
        logs = [log for log in logs if log.student_id == request.state.user]
    return logs

class BuyPagesBody(BaseModel):
    pages: int
@app.post("/buy_pages")
async def buy_pages(request: Request, body: BuyPagesBody):
    if auth.role(request.state.user) != "student":
        raise HTTPException(status_code=403, detail="Forbidden")
    student_id = request.state.user
    student = db.get_student_by_id(student_id)
    if student is None:
        raise HTTPException(status_code=404, detail="Student not found")
    student.remaining_pages += body.pages
    updated_student = db.update_student(student)
    return updated_student
